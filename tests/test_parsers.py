"""Parser tests."""

from docx import Document

from src.parsers.docx_parser import DocxParser


def test_docx_parser_converts_paragraphs_to_markdown(tmp_path):
    source = tmp_path / "example.docx"
    document = Document()
    document.add_heading("Title", level=1)
    document.add_paragraph("Paragraph text")
    document.save(source)

    markdown = DocxParser().parse(source)

    assert "# Title" in markdown
    assert "Paragraph text" in markdown
